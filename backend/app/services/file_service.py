import io
import os
import re
from typing import Optional, List, Dict, Any
from fastapi import UploadFile, HTTPException

# Import with fallbacks
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
except ImportError:
    Image = None

# Advanced vision model for handwriting recognition (primary)
try:
    from app.services.vision_model_service import vision_model_service
    VISION_MODEL_AVAILABLE = vision_model_service.is_available()
except ImportError:
    vision_model_service = None
    VISION_MODEL_AVAILABLE = False
except Exception:
    vision_model_service = None
    VISION_MODEL_AVAILABLE = False

# TrOCR service for handwriting recognition (fallback)
try:
    from app.services.trocr_service import trocr_service
    TROCR_AVAILABLE = trocr_service.is_available()
except ImportError:
    trocr_service = None
    TROCR_AVAILABLE = False

# Fallback to pytesseract if other models not available
try:
    import pytesseract
    # Configure pytesseract to find Tesseract on Windows
    import platform
    if platform.system() == 'Windows':
        # Common Windows installation paths
        tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        for path in tesseract_paths:
            import os
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                break
    TESSERACT_AVAILABLE = True
except ImportError:
    pytesseract = None
    TESSERACT_AVAILABLE = False
except Exception as e:
    pytesseract = None
    TESSERACT_AVAILABLE = False

class FileProcessingService:
    @staticmethod
    def _check_pdf_support() -> bool:
        """Check if PDF processing is available"""
        return pdfplumber is not None or PyPDF2 is not None
    
    @staticmethod
    def _check_docx_support() -> bool:
        """Check if DOCX processing is available"""
        return Document is not None
    
    @staticmethod
    def _check_image_support() -> bool:
        """Check if image processing is available (TrOCR preferred, tesseract fallback)"""
        return Image is not None and (TROCR_AVAILABLE or TESSERACT_AVAILABLE)
    
    @staticmethod
    async def extract_text_from_file(file: UploadFile) -> str:
        """Extract text from uploaded file based on file type"""
        try:
            content = await file.read()
            file_extension = file.filename.split('.')[-1].lower() if file.filename else ''
            
            if file_extension == 'pdf':
                if not FileProcessingService._check_pdf_support():
                    raise HTTPException(
                        status_code=500,
                        detail="PDF processing not available. Please install PyPDF2."
                    )
                return FileProcessingService._extract_from_pdf(content)
            elif file_extension in ['doc', 'docx']:
                if not FileProcessingService._check_docx_support():
                    raise HTTPException(
                        status_code=500,
                        detail="DOCX processing not available. Please install python-docx."
                    )
                return FileProcessingService._extract_from_docx(content)
            elif file_extension in ['jpg', 'jpeg', 'png']:
                if not FileProcessingService._check_image_support():
                    raise HTTPException(
                        status_code=500,
                        detail="Image OCR not available. TrOCR or pytesseract required."
                    )
                return FileProcessingService._extract_from_image(content)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file_extension}"
                )
                
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error processing file: {str(e)}"
            )

    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF file. Falls back to OCR for scanned/image PDFs."""
        if not pdfplumber and not PyPDF2:
            raise Exception("No PDF processing library available. Please install pdfplumber or PyPDF2.")
        
        text = ""
        
        try:
            # Try pdfplumber first (better text extraction)
            if pdfplumber:
                pdf_file = io.BytesIO(content)
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    return text.strip()
            
            # Fallback to PyPDF2 if pdfplumber fails or isn't available
            if PyPDF2:
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    return text.strip()
            
            # If no text found, try OCR on PDF pages (for scanned/image PDFs)
            ocr_text = FileProcessingService._extract_from_pdf_with_ocr(content)
            if ocr_text and ocr_text.strip():
                return ocr_text.strip()
            
            return "No readable text found in PDF"
            
        except Exception as e:
            raise Exception(f"Error extracting from PDF: {str(e)}")
    
    @staticmethod
    def _extract_from_pdf_with_ocr(content: bytes) -> str:
        """Extract text from PDF by converting pages to images and running OCR"""
        if not Image:
            return ""
        
        if not TROCR_AVAILABLE and not TESSERACT_AVAILABLE:
            return ""
        
        try:
            # Try pdf2image if available
            try:
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(content)
            except ImportError:
                # Try using pdfplumber to get page images
                if pdfplumber:
                    pdf_file = io.BytesIO(content)
                    images = []
                    with pdfplumber.open(pdf_file) as pdf:
                        for page in pdf.pages:
                            # Convert page to image
                            img = page.to_image(resolution=150)
                            images.append(img.original)
                else:
                    return ""
            
            if not images:
                return ""
            
            # Extract text from each page image
            all_text = []
            for i, img in enumerate(images):
                try:
                    page_text = FileProcessingService._ocr_image(img)
                    if page_text.strip():
                        all_text.append(f"--- Page {i+1} ---\n{page_text}")
                except Exception:
                    continue
            
            return "\n\n".join(all_text)
            
        except Exception as e:
            print(f"PDF OCR extraction failed: {e}")
            return ""
    
    @staticmethod
    def _ocr_image(image: Image.Image) -> str:
        """Run OCR on a PIL Image using advanced vision model with LLM reconstruction"""
        # Ensure RGB
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Use advanced vision model with LLM reconstruction (best for handwriting)
        if VISION_MODEL_AVAILABLE and vision_model_service:
            try:
                # Use full pipeline: OCR + LLM reconstruction
                raw_text, reconstructed_text = vision_model_service.extract_and_reconstruct(image)
                if reconstructed_text.strip():
                    return reconstructed_text
                elif raw_text.strip():
                    return raw_text
            except Exception as e:
                print(f"Vision model processing, trying fallback: {e}")
        
        # Fallback to TrOCR
        if TROCR_AVAILABLE and trocr_service:
            try:
                return trocr_service.extract_text_from_image_lines(image)
            except Exception as e:
                print(f"TrOCR failed: {e}")
        
        return ""

    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX file with improved formatting preservation"""
        if not Document:
            raise Exception("python-docx not available")
        
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs with better formatting
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    # Preserve line breaks and spacing for question formatting
                    text_parts.append(para_text)
            
            # Extract text from tables (often used for structured questions)
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.extend(table_text)
            
            if not text_parts:
                return "No readable text found in DOCX"
            
            # Join with newlines to preserve document structure
            return "\n".join(text_parts)
            
        except Exception as e:
            raise Exception(f"Error extracting from DOCX: {str(e)}")

    @staticmethod
    def parse_exam_structure(text: str) -> List[Dict[str, Any]]:
        """Parse exam text and extract structured question-answer pairs with enhanced pattern detection"""
        try:
            questions = []
            text = text.strip()
            
            # Enhanced question detection patterns
            question_patterns = [
                r'\n\s*(\d+)[\.):]\s*(.+?)(?=\n\s*\d+[\.):]\s*|\Z)',  # 1. Question text
                r'\n\s*([Qq]\d+)[\.):]\s*(.+?)(?=\n\s*[Qq]\d+[\.):]\s*|\Z)',  # Q1: Question text
                r'\n\s*([Qq]uestion\s+\d+)[\.):]?\s*(.+?)(?=\n\s*[Qq]uestion\s+\d+|\Z)',  # Question 1 text
                r'\n\s*(\d+\.\d+)\s*(.+?)(?=\n\s*\d+\.\d+\s*|\Z)',  # 1.1 Question text (sub-questions)
            ]
            
            sections = []
            for pattern in question_patterns:
                matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
                if matches:
                    sections = [(match[0], match[1].strip()) for match in matches]
                    break
            
            # If no clear pattern found, try simple number splitting
            if not sections:
                parts = re.split(r'\n\s*\d+[\.):]\s*', text)
                if len(parts) > 1:
                    # Skip first empty part and create numbered sections
                    for i, part in enumerate(parts[1:], 1):
                        if part.strip():
                            sections.append((str(i), part.strip()))
            
            # Process each section into structured question data
            for question_num, section_text in sections:
                if not section_text.strip():
                    continue
                
                question_data = FileProcessingService._parse_question_section_enhanced(
                    section_text.strip(), question_num
                )
                if question_data:
                    questions.append(question_data)
            
            # Fallback: treat entire text as one question if no structure found
            if not questions:
                questions = [{
                    "question_number": 1,
                    "question": "Please grade the following exam submission",
                    "student_answer": text,
                    "max_marks": 10,
                    "question_type": "essay"
                }]
            
            return questions
            
        except Exception as e:
            # Emergency fallback
            return [{
                "question_number": 1,
                "question": "Error parsing exam structure",
                "student_answer": text,
                "max_marks": 10,
                "question_type": "essay"
            }]
    
    @staticmethod
    def _parse_question_section(section: str, question_num: int) -> Dict[str, Any]:
        """Parse individual question section"""
        try:
            lines = section.split('\n')
            question_text = ""
            answer_text = ""
            question_type = "essay"
            max_marks = 5  # default
            
            # Look for answer patterns
            answer_patterns = [
                r'[Aa]nswer\s*:?\s*(.+)',
                r'[Ss]olution\s*:?\s*(.+)',
                r'[Rr]esponse\s*:?\s*(.+)'
            ]
            
            # Look for marks patterns
            marks_patterns = [
                r'\[(\d+)\s*marks?\]',
                r'\((\d+)\s*marks?\)',
                r'marks?\s*:?\s*(\d+)',
                r'points?\s*:?\s*(\d+)'
            ]
            
            # Detect multiple choice
            if re.search(r'[A-D][\.):]', section, re.IGNORECASE):
                question_type = "multiple_choice"
            
            current_section = "question"
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for marks
                for pattern in marks_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        max_marks = int(match.group(1))
                        line = re.sub(pattern, '', line, flags=re.IGNORECASE).strip()
                        break
                
                # Check if this line starts an answer
                is_answer = False
                for pattern in answer_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        current_section = "answer"
                        answer_text += match.group(1) + " "
                        is_answer = True
                        break
                
                if not is_answer:
                    if current_section == "question":
                        question_text += line + " "
                    else:
                        answer_text += line + " "
            
            # Clean up texts
            question_text = question_text.strip()
            answer_text = answer_text.strip()
            
            # If no clear question/answer separation, split roughly in half
            if not answer_text and question_text:
                parts = question_text.split('. ')
                if len(parts) > 2:
                    mid = len(parts) // 2
                    question_text = '. '.join(parts[:mid]) + '.'
                    answer_text = '. '.join(parts[mid:])
            
            return {
                "question_number": question_num,
                "question": question_text or f"Question {question_num}",
                "student_answer": answer_text or question_text,
                "max_marks": max_marks,
                "question_type": question_type
            }
            
        except Exception as e:
            return {
                "question_number": question_num,
                "question": f"Question {question_num}",
                "student_answer": section,
                "max_marks": 5,
                "question_type": "essay"
            }

    @staticmethod
    def _parse_question_section_enhanced(section: str, question_num: str) -> Dict[str, Any]:
        """Enhanced parsing of individual question section with better answer extraction"""
        try:
            lines = section.split('\n')
            question_text = ""
            answer_text = ""
            question_type = "essay"
            max_marks = 5  # default
            
            # Enhanced answer patterns
            answer_patterns = [
                r'[Aa]nswer\s*:?\s*(.+)',
                r'[Ss]olution\s*:?\s*(.+)',
                r'[Rr]esponse\s*:?\s*(.+)',
                r'[Aa]ns\s*:?\s*(.+)',
                r'[Ss]ol\s*:?\s*(.+)'
            ]
            
            # Enhanced marks patterns
            marks_patterns = [
                r'\[(\d+)\s*marks?\]',
                r'\((\d+)\s*marks?\)',
                r'marks?\s*:?\s*(\d+)',
                r'points?\s*:?\s*(\d+)',
                r'(\d+)\s*pts?',
                r'(\d+)\s*m\b'  # shorthand for marks
            ]
            
            # Enhanced question type detection
            if re.search(r'[A-D][\.):]', section, re.IGNORECASE):
                question_type = "multiple_choice"
            elif re.search(r'true|false', section, re.IGNORECASE):
                question_type = "true_false"
            elif re.search(r'fill\s+in\s+the\s+blank|complete\s+the\s+following', section, re.IGNORECASE):
                question_type = "fill_blank"
            elif re.search(r'define|explain|describe|discuss|analyze', section, re.IGNORECASE):
                question_type = "subjective"
            elif re.search(r'calculate|solve|find', section, re.IGNORECASE):
                question_type = "numeric"
            
            # Process lines more intelligently
            current_section = "question"
            question_lines = []
            answer_lines = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for marks first
                for pattern in marks_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        max_marks = int(match.group(1))
                        line = re.sub(pattern, '', line, flags=re.IGNORECASE).strip()
                        break
                
                # Check for answer indicators
                is_answer_line = False
                for pattern in answer_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        current_section = "answer"
                        answer_lines.append(match.group(1))
                        is_answer_line = True
                        break
                
                if not is_answer_line:
                    if current_section == "question":
                        question_lines.append(line)
                    else:
                        answer_lines.append(line)
            
            # Build final question and answer text
            question_text = ' '.join(question_lines).strip()
            answer_text = ' '.join(answer_lines).strip()
            
            # Smart fallback: if no clear separation, split by common patterns
            if not answer_text and question_text:
                # Look for question mark as separator
                if '?' in question_text:
                    parts = question_text.split('?', 1)
                    if len(parts) == 2 and parts[1].strip():
                        question_text = parts[0] + '?'
                        answer_text = parts[1].strip()
                # Look for colon as separator
                elif ':' in question_text and len(question_text.split(':')) == 2:
                    parts = question_text.split(':', 1)
                    question_text = parts[0].strip()
                    answer_text = parts[1].strip()
                # Split roughly in half as last resort
                elif len(question_text.split('.')) > 3:
                    sentences = question_text.split('.')
                    mid = len(sentences) // 2
                    question_text = '.'.join(sentences[:mid]) + '.'
                    answer_text = '.'.join(sentences[mid:])
            
            # Convert question_num to integer if possible
            try:
                match = re.search(r'\d+', str(question_num))
                q_num = int(match.group()) if match else 1
            except:
                q_num = 1
            
            return {
                "question_number": q_num,
                "question": question_text or f"Question {q_num}",
                "student_answer": answer_text or question_text or section,
                "max_marks": max_marks,
                "question_type": question_type
            }
            
        except Exception as e:
            # Fallback with error handling
            try:
                match = re.search(r'\d+', str(question_num))
                q_num = int(match.group()) if match else 1
            except:
                q_num = 1
                
            return {
                "question_number": q_num,
                "question": f"Question {q_num}",
                "student_answer": section,
                "max_marks": 5,
                "question_type": "essay"
            }

    @staticmethod
    def _extract_from_image(content: bytes) -> str:
        """Extract text from image using TrOCR (preferred) or Tesseract fallback"""
        if not Image:
            raise Exception("PIL not available")
        
        if not TROCR_AVAILABLE and not TESSERACT_AVAILABLE:
            raise Exception("No OCR engine available. Install transformers+torch for TrOCR or pytesseract.")
        
        try:
            image = Image.open(io.BytesIO(content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Use the shared OCR function
            text = FileProcessingService._ocr_image(image)
            
            if not text.strip():
                return "No readable text found in image"
            
            return text.strip()
            
        except Exception as e:
            return f"Could not extract text from image: {str(e)}"

    @staticmethod
    async def process_exam_file(file: UploadFile) -> Dict[str, Any]:
        """Process exam file and return structured data"""
        try:
            # Extract text from file
            extracted_text = await FileProcessingService.extract_text_from_file(file)
            
            # Parse exam structure
            parsed_questions = FileProcessingService.parse_exam_structure(extracted_text)
            
            return {
                "filename": file.filename,
                "extracted_text": extracted_text,
                "parsed_questions": parsed_questions,
                "total_questions": len(parsed_questions)
            }
            
        except Exception as e:
            raise Exception(f"Error processing exam file: {str(e)}")

    @staticmethod
    def validate_file_size(file_size: int, max_size_mb: int = 10) -> bool:
        """Validate file size"""
        max_size_bytes = max_size_mb * 1024 * 1024
        return file_size <= max_size_bytes

    @staticmethod
    def validate_file_type(filename: str) -> bool:
        """Validate file type - only allow specific formats for exam upload"""
        allowed_extensions = {'pdf', 'docx', 'jpg', 'jpeg', 'png'}
        
        if not filename:
            return False
            
        file_extension = filename.split('.')[-1].lower()
        return file_extension in allowed_extensions

# Create service instance
file_processing_service = FileProcessingService()